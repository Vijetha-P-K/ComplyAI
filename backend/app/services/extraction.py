import os

from docx import Document as DocxDocument
from PIL import Image
from pypdf import PdfReader

MAX_CHARS = 24000


def extract_text(path: str, file_type: str) -> str:
    if file_type == "pdf":
        text = _extract_pdf(path)
    elif file_type == "docx":
        text = _extract_docx(path)
    elif file_type in ("png", "jpg", "jpeg", "webp", "bmp", "tiff"):
        text = _extract_image(path)
    else:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    return text.strip()[:MAX_CHARS]


def _extract_pdf(path: str) -> str:
    reader = PdfReader(path)
    parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(parts)


def _extract_docx(path: str) -> str:
    doc = DocxDocument(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_image(path: str) -> str:
    import pytesseract

    image = Image.open(path)
    return pytesseract.image_to_string(image)


def get_file_type(filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower().lstrip(".")
    return ext or "txt"


ALLOWED_TYPES = {"pdf", "docx", "txt", "md", "csv", "png", "jpg", "jpeg", "webp", "bmp", "tiff"}
